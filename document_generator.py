"""
Модуль для генерации документов в формате Word (.docx)
ИСПРАВЛЕННАЯ ВЕРСИЯ с правильным mapping данных
"""

import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import re

class DocumentGenerator:
    def __init__(self, templates_dir="/home/ubuntu/templates"):
        self.templates_dir = templates_dir
        self.template_mapping = {
            'sok': 'sok_template.html',
            'instruction': 'instruction_template.html', 
            'procedure': 'procedure_template.html'
        }
    
    def generate_sok_docx(self, data):
        """
        Генерирует СОК в формате Word
        Принимает data словарь с ключами:
        - company_name, business_area, process_name, target_audience, 
        - process_steps, required_resources, expected_results
        """
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Заголовок
        title = doc.add_heading('СТАНДАРТНАЯ ОПЕРАЦИОННАЯ КАРТА (СОК)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Основная информация
        doc.add_paragraph()
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        # ИСПРАВЛЕННЫЙ MAPPING ДАННЫХ:
        info_data = [
            ('Компания:', data.get('company_name', 'Не указано')),
            ('Операция:', data.get('process_name', 'Не указано')),
            ('Исполнители:', data.get('target_audience', 'Не указано')),
            ('Дата создания:', data.get('creation_date', datetime.now().strftime('%d.%m.%Y')))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
        
        # Заголовок секции операций
        doc.add_paragraph()
        doc.add_heading('ПОСЛЕДОВАТЕЛЬНОСТЬ ОПЕРАЦИЙ', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Парсим шаги из process_steps
        steps_text = data.get('process_steps', '')
        if steps_text and steps_text != 'Не указано':
            # Разбиваем по цифрам или точкам
            steps = []
            if '1.' in steps_text or '2.' in steps_text:
                # Если есть нумерация
                steps = [step.strip() for step in re.split(r'\d+\.', steps_text) if step.strip()]
            else:
                # Если нет нумерации, разбиваем по предложениям
                steps = [step.strip() for step in steps_text.split('.') if step.strip()]
            
            if steps:
                ops_table = doc.add_table(rows=len(steps) + 1, cols=4)
                ops_table.style = 'Table Grid'
                ops_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Заголовки таблицы
                headers = ['№', 'Операция', 'Описание', 'Контроль']
                for i, header in enumerate(headers):
                    cell = ops_table.cell(0, i)
                    cell.text = header
                    # Жирный шрифт для заголовков
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Заполнение операций
                for i, step in enumerate(steps, 1):
                    ops_table.cell(i, 0).text = str(i)
                    ops_table.cell(i, 1).text = f"Этап {i}"
                    ops_table.cell(i, 2).text = step
                    ops_table.cell(i, 3).text = "✓"
        
        # Ресурсы и инструменты
        doc.add_paragraph()
        doc.add_heading('НЕОБХОДИМЫЕ РЕСУРСЫ', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        resources_text = data.get('required_resources', 'Не указано')
        doc.add_paragraph(f"Инструменты и материалы: {resources_text}")
        
        # Ожидаемые результаты
        doc.add_paragraph()
        doc.add_heading('ОЖИДАЕМЫЕ РЕЗУЛЬТАТЫ', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        results_text = data.get('expected_results', 'Не указано')
        doc.add_paragraph(f"Результат выполнения: {results_text}")
        
        # Согласование
        doc.add_paragraph()
        doc.add_heading('СОГЛАСОВАНИЕ', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        approval_table = doc.add_table(rows=2, cols=3)
        approval_table.style = 'Table Grid'
        
        # Заголовки
        headers = ['Разработал', 'Проверил', 'Утвердил']
        for i, header in enumerate(headers):
            cell = approval_table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Подписи
        signatures = [
            data.get('author', '_________________'),
            data.get('coordinator', '_________________'),
            data.get('approver', '_________________')
        ]
        
        for i, signature in enumerate(signatures):
            approval_table.cell(1, i).text = f"{signature}\n(подпись, дата)"
        
        # Подпись платформы
        doc.add_paragraph()
        footer = doc.add_paragraph('Создано с помощью платформы HowDo')
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return doc
    
    def generate_instruction_docx(self, data):
        """Генерирует рабочую инструкцию в формате Word"""
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Заголовок
        title = doc.add_heading('РАБОЧАЯ ИНСТРУКЦИЯ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Основная информация
        doc.add_paragraph()
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Компания:', data.get('company_name', 'Не указано')),
            ('Процесс:', data.get('process_name', 'Не указано')),
            ('Целевая аудитория:', data.get('target_audience', 'Не указано')),
            ('Версия:', data.get('version', '1.0')),
            ('Дата создания:', data.get('creation_date', datetime.now().strftime('%d.%m.%Y')))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
        
        # Цель и область применения
        doc.add_paragraph()
        doc.add_heading('ЦЕЛЬ И ОБЛАСТЬ ПРИМЕНЕНИЯ', 2)
        doc.add_paragraph(f"Сфера деятельности: {data.get('business_area', 'Не указано')}")
        doc.add_paragraph(f"Процесс: {data.get('process_name', 'Не указано')}")
        
        # Пошаговые инструкции
        doc.add_paragraph()
        doc.add_heading('ПОШАГОВЫЕ ИНСТРУКЦИИ', 2)
        steps_text = data.get('process_steps', 'Не указано')
        doc.add_paragraph(steps_text)
        
        # Необходимые ресурсы
        doc.add_paragraph()
        doc.add_heading('НЕОБХОДИМЫЕ РЕСУРСЫ', 2)
        doc.add_paragraph(data.get('required_resources', 'Не указано'))
        
        # Ожидаемые результаты
        doc.add_paragraph()
        doc.add_heading('ОЖИДАЕМЫЕ РЕЗУЛЬТАТЫ', 2)
        doc.add_paragraph(data.get('expected_results', 'Не указано'))
        
        # Согласование
        doc.add_paragraph()
        doc.add_heading('СОГЛАСОВАНИЕ', 2)
        
        approval_table = doc.add_table(rows=2, cols=3)
        approval_table.style = 'Table Grid'
        
        headers = ['Разработал', 'Проверил', 'Утвердил']
        for i, header in enumerate(headers):
            approval_table.cell(0, i).text = header
        
        signatures = [
            data.get('author', '_________________'),
            data.get('coordinator', '_________________'),
            data.get('approver', '_________________')
        ]
        
        for i, signature in enumerate(signatures):
            approval_table.cell(1, i).text = f"{signature}\n(подпись, дата)"
        
        # Подпись платформы
        doc.add_paragraph()
        footer = doc.add_paragraph('Создано с помощью платформы HowDo')
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return doc
    
    def generate_procedure_docx(self, data):
        """Генерирует стандарт процедуры в формате Word"""
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Заголовок
        title = doc.add_heading('СТАНДАРТ ПРОЦЕДУРЫ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Основная информация
        doc.add_paragraph()
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Компания:', data.get('company_name', 'Не указано')),
            ('Процедура:', data.get('process_name', 'Не указано')),
            ('Ответственные:', data.get('target_audience', 'Не указано')),
            ('Версия:', data.get('version', '1.0')),
            ('Дата создания:', data.get('creation_date', datetime.now().strftime('%d.%m.%Y')))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
        
        # Назначение процедуры
        doc.add_paragraph()
        doc.add_heading('НАЗНАЧЕНИЕ ПРОЦЕДУРЫ', 2)
        doc.add_paragraph(f"Область применения: {data.get('business_area', 'Не указано')}")
        
        # Описание процедуры
        doc.add_paragraph()
        doc.add_heading('ОПИСАНИЕ ПРОЦЕДУРЫ', 2)
        doc.add_paragraph(data.get('process_steps', 'Не указано'))
        
        # Ресурсы
        doc.add_paragraph()
        doc.add_heading('НЕОБХОДИМЫЕ РЕСУРСЫ', 2)
        doc.add_paragraph(data.get('required_resources', 'Не указано'))
        
        # Критерии качества
        doc.add_paragraph()
        doc.add_heading('КРИТЕРИИ КАЧЕСТВА', 2)
        doc.add_paragraph(data.get('expected_results', 'Не указано'))
        
        # Согласование
        doc.add_paragraph()
        doc.add_heading('СОГЛАСОВАНИЕ', 2)
        
        approval_table = doc.add_table(rows=2, cols=3)
        approval_table.style = 'Table Grid'
        
        headers = ['Разработал', 'Проверил', 'Утвердил']
        for i, header in enumerate(headers):
            approval_table.cell(0, i).text = header
        
        signatures = [
            data.get('author', '_________________'),
            data.get('coordinator', '_________________'),
            data.get('approver', '_________________')
        ]
        
        for i, signature in enumerate(signatures):
            approval_table.cell(1, i).text = f"{signature}\n(подпись, дата)"
        
        # Подпись платформы
        doc.add_paragraph()
        footer = doc.add_paragraph('Создано с помощью платформы HowDo')
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return doc
    
    # HTML методы для предварительного просмотра (упрощенные версии)
    def generate_sok_html(self, data):
        """Генерирует HTML для предварительного просмотра СОК"""
        html = f"""
        <html>
        <head>
            <meta charset="utf-8">
            <title>СОК - {data.get('process_name', 'Документ')}</title>
            <style>
                body {{ font-family: 'Times New Roman', serif; margin: 40px; }}
                h1 {{ text-align: center; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid black; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; font-weight: bold; }}
            </style>
        </head>
        <body>
            <h1>СТАНДАРТНАЯ ОПЕРАЦИОННАЯ КАРТА (СОК)</h1>
            
            <table>
                <tr><td><strong>Компания:</strong></td><td>{data.get('company_name', 'Не указано')}</td></tr>
                <tr><td><strong>Операция:</strong></td><td>{data.get('process_name', 'Не указано')}</td></tr>
                <tr><td><strong>Исполнители:</strong></td><td>{data.get('target_audience', 'Не указано')}</td></tr>
                <tr><td><strong>Дата создания:</strong></td><td>{data.get('creation_date', datetime.now().strftime('%d.%m.%Y'))}</td></tr>
            </table>
            
            <h2>ПОСЛЕДОВАТЕЛЬНОСТЬ ОПЕРАЦИЙ</h2>
            <p>{data.get('process_steps', 'Не указано')}</p>
            
            <h2>НЕОБХОДИМЫЕ РЕСУРСЫ</h2>
            <p>{data.get('required_resources', 'Не указано')}</p>
            
            <h2>ОЖИДАЕМЫЕ РЕЗУЛЬТАТЫ</h2>
            <p>{data.get('expected_results', 'Не указано')}</p>
            
            <p style="text-align: right; margin-top: 40px;"><em>Создано с помощью платформы HowDo</em></p>
        </body>
        </html>
        """
        return html
    
    def generate_instruction_html(self, data):
        """Генерирует HTML для предварительного просмотра инструкции"""
        return self.generate_sok_html(data).replace('СТАНДАРТНАЯ ОПЕРАЦИОННАЯ КАРТА (СОК)', 'РАБОЧАЯ ИНСТРУКЦИЯ')
    
    def generate_procedure_html(self, data):
        """Генерирует HTML для предварительного просмотра процедуры"""
        return self.generate_sok_html(data).replace('СТАНДАРТНАЯ ОПЕРАЦИОННАЯ КАРТА (СОК)', 'СТАНДАРТ ПРОЦЕДУРЫ')

