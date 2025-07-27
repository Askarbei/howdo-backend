from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import json
import os
from datetime import datetime
import uuid
from docx import Document
from docx.shared import Inches
import io

app = Flask(__name__)
CORS(app)

# Хранилище данных в памяти
users_db = {}
documents_db = {}

@app.route('/api/health')
def health():
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "users_count": len(users_db),
        "documents_count": len(documents_db)
    })

@app.route('/api/register', methods=['POST'])
def register():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    
    if not email or not password:
        return jsonify({"error": "Email и пароль обязательны"}), 400
    
    if email in users_db:
        return jsonify({"error": "Пользователь уже существует"}), 400
    
    user_id = str(uuid.uuid4())
    users_db[email] = {
        "id": user_id,
        "email": email,
        "password": password,
        "created_at": datetime.now().isoformat()
    }
    
    return jsonify({
        "message": "Пользователь зарегистрирован",
        "user": {"id": user_id, "email": email}
    })

@app.route('/api/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    
    if email not in users_db:
        return jsonify({"error": "Неверные учетные данные"}), 401
    
    user = users_db[email]
    if user["password"] != password:
        return jsonify({"error": "Неверные учетные данные"}), 401
    
    return jsonify({
        "message": "Успешный вход",
        "user": {"id": user["id"], "email": email}
    })

@app.route('/api/wizard', methods=['POST'])
def create_document_from_wizard():
    data = request.get_json()
    user_id = data.get('user_id', 'anonymous')
    answers = data.get('answers', {})
    
    # Создаем документ
    doc_id = str(uuid.uuid4())
    
    # Сохраняем в базу
    documents_db[doc_id] = {
        "id": doc_id,
        "user_id": user_id,
        "title": answers.get('q2', 'Стандартная операционная карта'),
        "answers": answers,
        "created_at": datetime.now().isoformat()
    }
    
    return jsonify({
        "message": "Стандарт создан успешно!",
        "document_id": doc_id
    })

@app.route('/api/documents')
def get_documents():
    user_id = request.args.get('user_id')
    
    user_documents = []
    for doc_id, doc in documents_db.items():
        if doc["user_id"] == user_id:
            user_documents.append({
                "id": doc_id,
                "title": doc["title"],
                "created_at": doc["created_at"]
            })
    
    return jsonify({"documents": user_documents})

@app.route('/api/documents/<doc_id>/download')
def download_document(doc_id):
    if doc_id not in documents_db:
        return jsonify({"error": "Документ не найден"}), 404
    
    doc_data = documents_db[doc_id]
    answers = doc_data["answers"]
    
    # Создаем Word документ с правильной структурой
    doc = Document()
    
    # Заголовок СОК
    title = f"СОК: {answers.get('q2', 'Не указано')}"
    doc.add_heading(title, 0)
    
    # Информация о документе
    doc.add_heading('Информация о документе', level=1)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    # Заполняем таблицу информации
    info_data = [
        ('Компания:', answers.get('q1', 'Не указано')),
        ('Операция:', answers.get('q2', 'Не указано')),
        ('Исполнители:', answers.get('q3', 'Не указано')),
        ('Дата создания:', datetime.now().strftime('%d.%m.%Y'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
    
    # Описание
    doc.add_heading('Описание', level=1)
    description = answers.get('q4', 'Стандартная процедура выполнения операции')
    doc.add_paragraph(description)
    
    # Шаги выполнения
    doc.add_heading('Шаги выполнения', level=1)
    steps = answers.get('q5', '').split('\n')
    for i, step in enumerate(steps, 1):
        if step.strip():
            doc.add_paragraph(f"{i}. {step.strip()}")
    
    # Требования безопасности
    doc.add_heading('Требования безопасности', level=1)
    safety = answers.get('q6', '').split('\n')
    for requirement in safety:
        if requirement.strip():
            doc.add_paragraph(requirement.strip(), style='List Bullet')
    
    # Контроль качества
    doc.add_heading('Контроль качества', level=1)
    quality = answers.get('q7', '').split('\n')
    for check in quality:
        if check.strip():
            doc.add_paragraph(check.strip(), style='List Bullet')
    
    # Ожидаемые результаты
    if answers.get('q8'):
        doc.add_heading('Ожидаемые результаты', level=1)
        doc.add_paragraph(answers.get('q8', ''))
    
    # Сохраняем в память
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    filename = f"{answers.get('q2', 'document').replace(' ', '_')}.docx"
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)

